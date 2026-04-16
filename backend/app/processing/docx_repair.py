"""DOCX document repair: auto-fix detected issues.

Fixers are based on the thesis format standard defined in
docs/references/thesis-format-standard.md (Chinese undergraduate thesis).
"""

from __future__ import annotations

import io
import logging
import re
import unicodedata
from collections import Counter
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Emu
from lxml import etree

logger = logging.getLogger(__name__)

# OOXML schema child order for w:pPr (CT_PPr) and w:rPr (CT_RPr).
# Word/LibreOffice require children to appear in this order; inserting
# out-of-order elements can be ignored or cause the document to be
# rejected on open.
_PPR_CHILD_ORDER = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs",
    "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment",
    "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr",
    "pPrChange",
)

_RPR_CHILD_ORDER = (
    "rStyle", "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps",
    "strike", "dstrike", "outline", "shadow", "emboss", "imprint", "noProof",
    "snapToGrid", "vanish", "webHidden", "color", "spacing", "w", "kern",
    "position", "sz", "szCs", "highlight", "u", "effect", "bdr", "shd",
    "fitText", "vertAlign", "rtl", "cs", "em", "lang", "eastAsianLayout",
    "specVanish", "oMath",
)


def _local_name(tag: str) -> str:
    return tag.split("}", 1)[-1] if "}" in tag else tag


def _insert_in_schema_order(parent, child, order: tuple) -> None:
    """Insert child into parent at the schema-correct position."""
    child_local = _local_name(child.tag)
    if child_local not in order:
        parent.append(child)
        return
    target = order.index(child_local)
    insert_idx = len(parent)
    for i, existing in enumerate(parent):
        existing_local = _local_name(existing.tag)
        if existing_local in order and order.index(existing_local) <= target:
            continue
        insert_idx = i
        break
    parent.insert(insert_idx, child)


def _ensure_ppr_child(ppr, qname: str):
    """Return existing pPr child or create it at schema-correct position."""
    child = ppr.find(qname)
    if child is None:
        child = etree.Element(qname)
        _insert_in_schema_order(ppr, child, _PPR_CHILD_ORDER)
    return child


def _ensure_rpr_child(rpr, qname: str):
    """Return existing rPr child or create it at schema-correct position."""
    child = rpr.find(qname)
    if child is None:
        child = etree.Element(qname)
        _insert_in_schema_order(rpr, child, _RPR_CHILD_ORDER)
    return child


def _ensure_pPr(para):
    """Return existing pPr or create it as the first child of w:p."""
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = etree.Element(qn("w:pPr"))
        para._element.insert(0, ppr)  # pPr MUST be first child of w:p
    return ppr


def _ensure_rPr(run):
    """Return existing rPr or create it as the first child of w:r."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.Element(qn("w:rPr"))
        run._element.insert(0, rpr)  # rPr MUST be first child of w:r
    return rpr

# Standard values (must match docx_analyze.py)
STD_BODY_CJK_FONT = "宋体"
STD_BODY_LATIN_FONT = "Times New Roman"
STD_BODY_SIZE_PT = 12.0
STD_H1_SIZE_PT = 15.0
STD_H2_SIZE_PT = 14.0
STD_H3_SIZE_PT = 12.0
STD_MARGIN_LR_TWIPS = 1803
STD_MARGIN_TB_TWIPS = 1440


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def repair_docx(docx_bytes: bytes, issue_codes: list[str]) -> bytes:
    """Apply selected fixes to a DOCX document.

    Fixers are applied in a predefined order to guarantee dependency correctness.
    Supports synthetic codes like CITATION_DELETE_24 to remove specific citations.
    """
    doc = Document(io.BytesIO(docx_bytes))

    applied = []
    for code in _FIXER_ORDER:
        if code not in issue_codes:
            continue
        fixer = _FIXERS.get(code)
        if fixer is None:
            continue
        try:
            fixer(doc)
            applied.append(code)
        except Exception:
            logger.warning("Fixer %s failed, skipping", code, exc_info=True)

    # Handle CITATION_DELETE_N synthetic codes
    delete_nums = []
    for code in issue_codes:
        if code.startswith("CITATION_DELETE_"):
            try:
                num = int(code.split("_")[-1])
                delete_nums.append(num)
            except ValueError:
                pass
    if delete_nums:
        try:
            _fix_citation_delete(doc, delete_nums)
            applied.append(f"CITATION_DELETE:{delete_nums}")
        except Exception:
            logger.warning("Citation delete failed", exc_info=True)

    logger.info("Applied %d fixers: %s", len(applied), applied)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Thesis-standard fixers
# ---------------------------------------------------------------------------

def _fix_body_font(doc: Document) -> None:
    """FT-001/002: Normalize body text to SimSun (CJK) + Times New Roman (Latin).

    Sets both the ASCII/hAnsi font (via python-docx) and the eastAsia font
    (via direct XML manipulation) on each body text run.
    """
    changed = 0
    for para in _body_paras(doc):
        for run in para.runs:
            if not run.text.strip():
                continue
            modified = False
            # Set Latin font
            if run.font.name != STD_BODY_LATIN_FONT:
                run.font.name = STD_BODY_LATIN_FONT
                modified = True
            # Set CJK font via rFonts XML (schema-ordered insertion).
            rpr = _ensure_rPr(run)
            rfonts = _ensure_rpr_child(rpr, qn("w:rFonts"))
            if rfonts.get(qn("w:eastAsia")) != STD_BODY_CJK_FONT:
                rfonts.set(qn("w:eastAsia"), STD_BODY_CJK_FONT)
                modified = True
            # Also set hAnsi to match Latin
            if rfonts.get(qn("w:hAnsi")) != STD_BODY_LATIN_FONT:
                rfonts.set(qn("w:hAnsi"), STD_BODY_LATIN_FONT)
            if modified:
                changed += 1

    logger.debug("Normalized %d body runs to %s + %s", changed, STD_BODY_CJK_FONT, STD_BODY_LATIN_FONT)


def _fix_body_cjk_font(doc: Document) -> None:
    """FT-001 only: fix CJK font. Delegates to _fix_body_font."""
    _fix_body_font(doc)


def _fix_body_size(doc: Document) -> None:
    """FS-001: Set body text font size to 12pt (小四)."""
    target = Pt(STD_BODY_SIZE_PT)
    changed = 0
    for para in _body_paras(doc):
        for run in para.runs:
            if not run.text.strip():
                continue
            if run.font.size and abs(run.font.size.pt - STD_BODY_SIZE_PT) > 0.5:
                run.font.size = target
                changed += 1
    logger.debug("Set %d body runs to %spt", changed, STD_BODY_SIZE_PT)


def _fix_body_spacing(doc: Document) -> None:
    """SP-001: Set body line spacing to 1.5x (line=360, lineRule=auto)."""
    changed = 0
    for para in _body_paras(doc):
        ppr = _ensure_pPr(para)
        spacing = _ensure_ppr_child(ppr, qn("w:spacing"))

        cur_line = spacing.get(qn("w:line"))
        cur_rule = spacing.get(qn("w:lineRule"))

        needs_fix = False
        if cur_line != "360":
            needs_fix = True
        if cur_rule is not None and cur_rule != "auto":
            needs_fix = True

        if needs_fix:
            spacing.set(qn("w:line"), "360")
            spacing.set(qn("w:lineRule"), "auto")
            changed += 1

    logger.debug("Set %d body paragraphs to 1.5x line spacing", changed)


def _fix_body_indent(doc: Document) -> None:
    """SP-002: Set body first-line indent to 2 chars (firstLineChars=200)."""
    changed = 0
    for para in _body_paras(doc):
        ppr = _ensure_pPr(para)
        ind = _ensure_ppr_child(ppr, qn("w:ind"))

        cur = ind.get(qn("w:firstLineChars"))
        if cur != "200":
            ind.set(qn("w:firstLineChars"), "200")
            # Remove EMU-based firstLine if present (chars takes precedence)
            if ind.get(qn("w:firstLine")) is not None:
                del ind.attrib[qn("w:firstLine")]
            changed += 1

    logger.debug("Set %d body paragraphs to 2-char first-line indent", changed)


def _fix_body_alignment(doc: Document) -> None:
    """SP-007: Set body text alignment to justify."""
    changed = 0
    for para in _body_paras(doc):
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            changed += 1
    logger.debug("Set %d body paragraphs to justify alignment", changed)


def _fix_heading_size(doc: Document) -> None:
    """FS-002/003/004: Set heading font sizes to standard values."""
    expected = {1: Pt(STD_H1_SIZE_PT), 2: Pt(STD_H2_SIZE_PT), 3: Pt(STD_H3_SIZE_PT)}
    changed = 0
    for para in doc.paragraphs:
        level = _get_heading_level(para)
        if level is None or level not in expected:
            continue
        for run in para.runs:
            if not run.text.strip():
                continue
            if run.font.size and abs(run.font.size.pt - expected[level].pt) > 0.5:
                run.font.size = expected[level]
                changed += 1
    logger.debug("Normalized %d heading runs to standard sizes", changed)


def _fix_heading_alignment(doc: Document) -> None:
    """SP-003: Set H1 alignment to center."""
    changed = 0
    for para in doc.paragraphs:
        level = _get_heading_level(para)
        if level != 1:
            continue
        if para.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            changed += 1
    logger.debug("Set %d H1 headings to center alignment", changed)


def _fix_page_margins(doc: Document) -> None:
    """PL-002/003: Adjust page margins to standard values."""
    # Twips to EMU: 1 twip = 635 EMU (= 914400/1440)
    lr_emu = STD_MARGIN_LR_TWIPS * 635
    tb_emu = STD_MARGIN_TB_TWIPS * 635

    changed = 0
    for section in doc.sections:
        if section.left_margin != lr_emu:
            section.left_margin = lr_emu
            changed += 1
        if section.right_margin != lr_emu:
            section.right_margin = lr_emu
            changed += 1
        if section.top_margin != tb_emu:
            section.top_margin = tb_emu
            changed += 1
        if section.bottom_margin != tb_emu:
            section.bottom_margin = tb_emu
            changed += 1

    logger.debug("Adjusted margins on %d section properties", changed)


# ---------------------------------------------------------------------------
# Generic fixers (kept)
# ---------------------------------------------------------------------------

def _fix_redundant_empty_paragraphs(doc: Document) -> None:
    consecutive = 0
    to_remove = []
    for para in doc.paragraphs:
        if not para.text.strip():
            consecutive += 1
            if consecutive >= 2:
                to_remove.append(para._element)
        else:
            consecutive = 0
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    logger.debug("Removed %d redundant empty paragraphs", len(to_remove))


def _fix_missing_heading_styles(doc: Document) -> None:
    """Apply Heading styles to paragraphs matching Chinese numbering patterns."""
    level_map = _build_heading_level_map(doc)
    if not level_map:
        return

    changed = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) > 100 or _is_heading(para):
            continue
        level = _match_heading_pattern(text)
        if level is None:
            continue
        runs = [r for r in para.runs if r.text.strip()]
        if not runs or not all(r.bold for r in runs):
            continue

        heading_style = _find_or_create_heading_style(doc, level)
        if heading_style is None:
            continue
        para.style = heading_style
        changed += 1

    logger.debug("Applied heading styles to %d paragraphs", changed)


def _fix_heading_level_flat(doc: Document) -> None:
    """Re-level existing headings based on Chinese numbering patterns.

    When all headings are at the same level (e.g., all H1) but their text
    uses numbering patterns that imply different levels (一、 = H1,
    （一）= H2, 1. = H3), this fixer reassigns the correct heading level.

    Headings without a recognized pattern keep their current level.
    """
    changed = 0
    for para in doc.paragraphs:
        if not _is_heading(para):
            continue
        text = para.text.strip()
        if not text:
            continue

        target_level = _match_heading_pattern(text)
        if target_level is None:
            continue

        current_level = _get_heading_level(para)
        if current_level == target_level:
            continue

        # Apply correct heading style
        style = _find_or_create_heading_style(doc, target_level)
        if style is None:
            continue
        para.style = style
        changed += 1

    logger.debug("Re-leveled %d headings based on numbering patterns", changed)


def _fix_heading_orphan(doc: Document) -> None:
    changed = 0
    for para in doc.paragraphs:
        if _is_heading(para) and not para.paragraph_format.keep_with_next:
            para.paragraph_format.keep_with_next = True
            changed += 1
    logger.debug("Set keep_with_next on %d headings", changed)


def _fix_toc_outdated(doc: Document) -> None:
    body = doc.element.body
    for sdt in body.iter(qn("w:sdt")):
        sdt_pr = sdt.find(qn("w:sdtPr"))
        if sdt_pr is None:
            continue
        doc_part = sdt_pr.find(qn("w:docPartObj"))
        if doc_part is None:
            continue
        gallery = doc_part.find(qn("w:docPartGallery"))
        if gallery is not None and "Table of Contents" in (gallery.get(qn("w:val")) or ""):
            _mark_fields_dirty(sdt)
            return
    for fld in body.iter(qn("w:fldChar")):
        if fld.get(qn("w:fldCharType")) == "begin":
            parent = fld.getparent()
            if parent is not None:
                for sibling in parent.itersiblings():
                    instr = sibling.find(qn("w:instrText"))
                    if instr is not None and instr.text and "TOC" in instr.text:
                        fld.set(qn("w:dirty"), "true")
                        return
                    break


def _fix_citation_delete(doc: Document, numbers: list[int]) -> None:
    """Remove specific [N] citation markers from body text.

    Citations like [24] may be split across multiple runs (e.g., run="[",
    run="24", run="]"). This fixer works at the paragraph level: it
    concatenates run texts, performs the deletion on the combined string,
    then redistributes the result back to the runs.
    """
    if not numbers:
        return

    num_set = set(numbers)
    cite_re = re.compile(r"\[(\d+(?:[,，\-]\d+)*)\]")

    changed = 0
    for para in doc.paragraphs:
        if _is_heading(para):
            continue

        full_text = para.text
        if not full_text or not any(f"[{n}]" in full_text or f"{n}" in full_text for n in num_set):
            continue

        # Build a mapping: character position in full_text -> (run_index, char_offset_in_run)
        runs = para.runs
        if not runs:
            continue

        # Concatenate run texts and find citation matches on the full string
        run_texts = [r.text or "" for r in runs]
        combined = "".join(run_texts)

        # Find all citation brackets to delete
        deletions: list[tuple[int, int]] = []  # (start, end) in combined string
        for m in cite_re.finditer(combined):
            raw = m.group(1)
            cited = _parse_cite_numbers(raw)
            if not cited:
                continue
            if cited.issubset(num_set):
                # Delete entire bracket
                deletions.append((m.start(), m.end()))
            elif cited & num_set:
                # Partial: keep remaining numbers
                remaining = sorted(cited - num_set)
                new_bracket = "[" + ",".join(str(n) for n in remaining) + "]"
                deletions.append((m.start(), m.end(), new_bracket))  # type: ignore[arg-type]

        if not deletions:
            continue

        # Apply deletions in reverse order to preserve positions
        result = combined
        for d in reversed(deletions):
            if len(d) == 3:
                start, end, replacement = d  # type: ignore[misc]
                result = result[:start] + replacement + result[end:]
            else:
                start, end = d
                result = result[:start] + result[end:]

        # Redistribute the result back to runs
        _redistribute_text(runs, result)
        changed += 1

    logger.debug("Removed citation references for %s from %d paragraphs", numbers, changed)


def _redistribute_text(runs, new_text: str) -> None:
    """Distribute new_text across existing runs, preserving formatting.

    Uses a character-budget approach: each run gets characters proportional
    to its original length. This keeps formatting boundaries roughly intact.
    """
    if not runs:
        return

    old_lengths = [len(r.text or "") for r in runs]
    old_total = sum(old_lengths)

    if old_total == 0:
        runs[0].text = new_text
        return

    pos = 0
    for i, run in enumerate(runs):
        if i == len(runs) - 1:
            # Last run gets remaining text
            run.text = new_text[pos:]
        else:
            # Proportional share
            share = round(len(new_text) * old_lengths[i] / old_total) if old_total > 0 else 0
            run.text = new_text[pos:pos + share]
            pos += share


def _parse_cite_numbers(raw: str) -> set[int]:
    """Parse citation number string like '1,2,5-8' into a set of ints."""
    nums: set[int] = set()
    for part in re.split(r"[,，]", raw):
        part = part.strip()
        if "-" in part:
            try:
                a, b = part.split("-", 1)
                for n in range(int(a), int(b) + 1):
                    nums.add(n)
            except ValueError:
                pass
        else:
            try:
                nums.add(int(part))
            except ValueError:
                pass
    return nums


def _fix_citation_ref_numbering(doc: Document) -> None:
    """Add [N] numbering prefix to reference entries that lack it.

    Locates the reference section (参考文献), then for each entry paragraph
    that doesn't already start with [N], prepends [N] to the first run.
    """
    ref_paras = _locate_ref_paragraphs(doc)
    if not ref_paras:
        logger.debug("No reference section found, skipping numbering fix")
        return

    import re
    prefix_re = re.compile(r"^\[\d+\]")

    changed = 0
    for idx, para in enumerate(ref_paras, start=1):
        text = para.text.strip()
        if not text:
            continue
        # Skip if already has [N] prefix
        if prefix_re.match(text):
            continue
        # Prepend [N] to the first run
        runs = para.runs
        if runs:
            runs[0].text = f"[{idx}]{runs[0].text}"
        else:
            # No runs (rare), add text directly
            para.text = f"[{idx}]{text}"
        changed += 1

    logger.debug("Added [N] prefix to %d reference entries", changed)


def _fix_heading_format_inconsistency(doc: Document) -> None:
    """Normalize heading format within each level.

    Sets standard font sizes and clears direct font-name overrides so that
    headings at the same level look consistent. Bold is preserved since
    the standard requires headings to be bold.
    """
    expected = {1: Pt(STD_H1_SIZE_PT), 2: Pt(STD_H2_SIZE_PT), 3: Pt(STD_H3_SIZE_PT)}
    changed = 0
    for para in doc.paragraphs:
        level = _get_heading_level(para)
        if level is None or level not in expected:
            continue
        for run in para.runs:
            if not run.text.strip():
                continue
            modified = False
            # Normalize size to standard
            if run.font.size and run.font.size != expected[level]:
                run.font.size = expected[level]
                modified = True
            # Clear direct font-name override (let style definition apply)
            if run.font.name is not None:
                run.font.name = None
                modified = True
            if modified:
                changed += 1
    logger.debug("Normalized %d heading runs", changed)


def _fix_image_overflow(doc: Document) -> None:
    try:
        section = doc.sections[0]
        pw = section.page_width - section.left_margin - section.right_margin
    except (IndexError, TypeError):
        return
    changed = 0
    for shape in doc.inline_shapes:
        if shape.width and shape.width > pw:
            ratio = pw / shape.width
            shape.width = pw
            if shape.height:
                shape.height = int(shape.height * ratio)
            changed += 1
    logger.debug("Resized %d overflowing images", changed)


# ---------------------------------------------------------------------------
# Fixer registry
# ---------------------------------------------------------------------------

_FIXER_ORDER = [
    # Thesis-standard fixers (body text)
    "BODY_FONT_WRONG",
    "BODY_CJK_FONT_WRONG",
    "BODY_SIZE_WRONG",
    "BODY_SPACING_WRONG",
    "BODY_INDENT_WRONG",
    "BODY_ALIGNMENT_WRONG",
    # Thesis-standard fixers (headings)
    "HEADING_SIZE_WRONG",
    "HEADING_ALIGNMENT_WRONG",
    # Thesis-standard fixers (page)
    "PAGE_MARGINS_WRONG",
    # Citation fixers
    "CITATION_REF_NUMBERING_MISSING",
    # Generic fixers
    "REDUNDANT_EMPTY_PARAGRAPHS",
    "MISSING_HEADING_STYLES",
    "HEADING_LEVEL_FLAT",
    "HEADING_ORPHAN",
    "TOC_OUTDATED",
    "HEADING_FORMAT_INCONSISTENCY",
    "IMAGE_OVERFLOW",
]

_FIXERS: dict[str, Any] = {
    "BODY_FONT_WRONG": _fix_body_font,
    "BODY_CJK_FONT_WRONG": _fix_body_cjk_font,
    "BODY_SIZE_WRONG": _fix_body_size,
    "BODY_SPACING_WRONG": _fix_body_spacing,
    "BODY_INDENT_WRONG": _fix_body_indent,
    "BODY_ALIGNMENT_WRONG": _fix_body_alignment,
    "HEADING_SIZE_WRONG": _fix_heading_size,
    "HEADING_ALIGNMENT_WRONG": _fix_heading_alignment,
    "PAGE_MARGINS_WRONG": _fix_page_margins,
    "CITATION_REF_NUMBERING_MISSING": _fix_citation_ref_numbering,
    "REDUNDANT_EMPTY_PARAGRAPHS": _fix_redundant_empty_paragraphs,
    "MISSING_HEADING_STYLES": _fix_missing_heading_styles,
    "HEADING_LEVEL_FLAT": _fix_heading_level_flat,
    "HEADING_ORPHAN": _fix_heading_orphan,
    "TOC_OUTDATED": _fix_toc_outdated,
    "HEADING_FORMAT_INCONSISTENCY": _fix_heading_format_inconsistency,
    "IMAGE_OVERFLOW": _fix_image_overflow,
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


def _is_heading(para) -> bool:
    return bool(para.style and para.style.name and para.style.name.startswith("Heading"))


def _get_heading_level(para) -> int | None:
    if not para.style or not para.style.name:
        return None
    m = _HEADING_RE.match(para.style.name)
    return int(m.group(1)) if m else None


def _body_paras(doc: Document):
    """Yield non-empty, non-heading paragraphs."""
    for para in doc.paragraphs:
        if _is_heading(para) or not para.text.strip():
            continue
        yield para


# -- Heading pattern matching --

_HEADING_PATTERNS: list[tuple[re.Pattern, int]] = [
    (re.compile(r'^[一二三四五六七八九十]+、\S'), 1),
    (re.compile(r'^[（(][一二三四五六七八九十]+[)）]\S'), 2),
    (re.compile(r'^(Abstract|摘\s*要|参考文献|致\s*谢|附\s*录)$'), 1),
]

_HEADING_MAX_LEN = 40


def _match_heading_pattern(text: str) -> int | None:
    text = text.strip()
    if len(text) > _HEADING_MAX_LEN:
        return None
    for pattern, level in _HEADING_PATTERNS:
        if pattern.search(text):
            return level
    return None


def _build_heading_level_map(doc: Document) -> dict[str, int]:
    counts: dict[str, int] = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) > 100:
            continue
        level = _match_heading_pattern(text)
        if level is not None:
            key = f"level_{level}"
            counts[key] = counts.get(key, 0) + 1
    return counts


def _find_or_create_heading_style(doc: Document, level: int):
    name = f"Heading {level}"
    try:
        return doc.styles[name]
    except KeyError:
        pass
    try:
        from docx.enum.style import WD_STYLE_TYPE
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = None
        for existing_level in [1, 2, 3, 4]:
            try:
                style.base_style = doc.styles[f"Heading {existing_level}"]
                break
            except KeyError:
                continue
        return style
    except Exception:
        return None


def _locate_ref_paragraphs(doc: Document) -> list:
    """Find paragraph objects in the reference section."""
    _REF_PATTERNS = [
        re.compile(r"^参\s*考\s*文\s*献$"),
        re.compile(r"^References?$", re.IGNORECASE),
    ]
    _STOP_PATTERNS = [
        re.compile(r"^致\s*谢"),
        re.compile(r"^Acknowledg", re.IGNORECASE),
        re.compile(r"^附\s*录"),
        re.compile(r"^Appendi", re.IGNORECASE),
    ]

    paras = doc.paragraphs
    start_idx = None
    for i, para in enumerate(paras):
        text = para.text.strip()
        if text and any(p.match(text) for p in _REF_PATTERNS):
            start_idx = i + 1
            break

    if start_idx is None:
        return []

    result = []
    for i in range(start_idx, len(paras)):
        para = paras[i]
        text = para.text.strip()
        if not text:
            continue
        # Stop at next section heading
        if _is_heading(para):
            break
        if any(p.match(text) for p in _STOP_PATTERNS):
            break
        result.append(para)

    return result


def _mark_fields_dirty(element) -> None:
    for fld in element.iter(qn("w:fldChar")):
        fld.set(qn("w:dirty"), "true")
