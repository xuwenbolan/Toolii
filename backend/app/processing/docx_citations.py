"""DOCX citation analysis: detect citation style and cross-check references."""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from typing import Any, Literal

from docx import Document


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def check_citations(doc: Document) -> list[dict[str, Any]]:
    """Run citation analysis and return 0-N issues.

    Two-phase pipeline:
      1. Identify citation style (numbered, author_year, footnote, mixed, none)
      2. Run style-specific checkers
    """
    profile = _build_citation_profile(doc)

    if profile.style == "none":
        return []

    issues: list[dict[str, Any]] = []
    for checker in _UNIVERSAL_CHECKERS:
        result = checker(doc, profile)
        if result is not None:
            issues.append(result)

    for checker in _STYLE_CHECKERS.get(profile.style, []):
        result = checker(doc, profile)
        if result is not None:
            issues.append(result)

    return issues


# ---------------------------------------------------------------------------
# Citation profile
# ---------------------------------------------------------------------------

@dataclass
class CitationProfile:
    style: Literal["numbered", "author_year", "footnote", "mixed", "none"]
    # Numbered style
    inline_numbers: set[int] = field(default_factory=set)
    ref_entry_count: int = 0
    ref_has_numbering: bool = False
    # Author-year style
    inline_author_year_count: int = 0
    # Footnote/endnote
    footnote_count: int = 0
    endnote_count: int = 0
    # Reference section
    has_ref_section: bool = False
    ref_entries: list[str] = field(default_factory=list)
    # Stats
    numbered_hit_count: int = 0


def _build_citation_profile(doc: Document) -> CitationProfile:
    """Scan document and build a citation profile."""
    # Scan inline citations
    inline_numbers, numbered_hits = _scan_numbered_citations(doc)
    author_year_count = _scan_author_year_citations(doc)

    # Scan footnotes/endnotes
    footnote_count = _count_footnotes(doc)
    endnote_count = _count_endnotes(doc)

    # Locate reference section
    ref_entries = _locate_reference_entries(doc)
    has_ref_section = len(ref_entries) > 0
    ref_has_numbering = _check_ref_numbering(ref_entries)

    # Determine dominant style
    style = _determine_style(
        numbered_hits, author_year_count,
        footnote_count + endnote_count,
    )

    return CitationProfile(
        style=style,
        inline_numbers=inline_numbers,
        ref_entry_count=len(ref_entries),
        ref_has_numbering=ref_has_numbering,
        inline_author_year_count=author_year_count,
        footnote_count=footnote_count,
        endnote_count=endnote_count,
        has_ref_section=has_ref_section,
        ref_entries=ref_entries,
        numbered_hit_count=numbered_hits,
    )


# ---------------------------------------------------------------------------
# Phase 1: Scanning
# ---------------------------------------------------------------------------

_BRACKET_CITE_RE = re.compile(r"\[(\d+(?:[,，\-]\d+)*)\]")


def _scan_numbered_citations(doc: Document) -> tuple[set[int], int]:
    """Scan body text for [N] style citations.

    Returns (set of cited numbers, total hit count).
    Excludes paragraphs in TOC styles.
    """
    cited: set[int] = set()
    total_hits = 0

    for para in doc.paragraphs:
        # Skip TOC entries (they contain page number brackets like [1])
        if para.style and para.style.name and para.style.name.startswith("toc"):
            continue
        text = para.text
        for m in _BRACKET_CITE_RE.finditer(text):
            total_hits += 1
            raw = m.group(1)
            cited.update(_parse_number_list(raw))

    return cited, total_hits


_AUTHOR_YEAR_CN_RE = re.compile(
    r"[（(]([^）)]{2,80}?(?:19|20)\d{2}[a-z]?)[）)]"
)
_AUTHOR_YEAR_EN_RE = re.compile(
    r"\(([A-Z][a-z]+[^)]{0,80}?(?:19|20)\d{2}[a-z]?)\)"
)


def _scan_author_year_citations(doc: Document) -> int:
    """Count author-year style citations in body text."""
    count = 0
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("toc"):
            continue
        text = para.text
        count += len(_AUTHOR_YEAR_CN_RE.findall(text))
        count += len(_AUTHOR_YEAR_EN_RE.findall(text))
    return count


def _count_footnotes(doc: Document) -> int:
    """Count Word footnotes (excluding default separator)."""
    try:
        footnotes_part = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        )
        if footnotes_part is None:
            return 0
        from lxml import etree
        from docx.oxml.ns import qn
        tree = etree.fromstring(footnotes_part.blob)
        # Footnote IDs 0 and 1 are separator/continuation, real ones start at 2
        footnotes = tree.findall(qn("w:footnote"))
        return max(0, len(footnotes) - 2)
    except Exception:
        return 0


def _count_endnotes(doc: Document) -> int:
    """Count Word endnotes (excluding default separator)."""
    try:
        endnotes_part = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
        )
        if endnotes_part is None:
            return 0
        from lxml import etree
        from docx.oxml.ns import qn
        tree = etree.fromstring(endnotes_part.blob)
        endnotes = tree.findall(qn("w:endnote"))
        return max(0, len(endnotes) - 2)
    except Exception:
        return 0


# ---------------------------------------------------------------------------
# Phase 1: Reference section locating
# ---------------------------------------------------------------------------

_REF_SECTION_PATTERNS = [
    re.compile(r"^参\s*考\s*文\s*献$"),
    re.compile(r"^References?$", re.IGNORECASE),
    re.compile(r"^Bibliography$", re.IGNORECASE),
    re.compile(r"^Works?\s+Cited$", re.IGNORECASE),
]

_STOP_PATTERNS = [
    re.compile(r"^致\s*谢"),
    re.compile(r"^Acknowledg", re.IGNORECASE),
    re.compile(r"^附\s*录"),
    re.compile(r"^Appendi", re.IGNORECASE),
]


def _is_ref_section_heading(para) -> bool:
    """Check if paragraph is a reference section heading."""
    text = para.text.strip()
    if not text:
        return False
    return any(p.match(text) for p in _REF_SECTION_PATTERNS)


def _is_stop_heading(para) -> bool:
    """Check if paragraph marks the end of the reference section."""
    text = para.text.strip()
    if not text:
        return False
    # Actual heading style
    if para.style and para.style.name and para.style.name.startswith("Heading"):
        return True
    # Pseudo-heading that is a known stop pattern
    return any(p.match(text) for p in _STOP_PATTERNS)


def _locate_reference_entries(doc: Document) -> list[str]:
    """Find reference section and extract non-empty entry texts."""
    paras = doc.paragraphs
    start_idx = None

    # 1. Search for reference heading (by style or by text)
    for i, para in enumerate(paras):
        if _is_ref_section_heading(para):
            start_idx = i + 1
            break

    if start_idx is None:
        return []

    entries: list[str] = []
    for i in range(start_idx, len(paras)):
        para = paras[i]
        text = para.text.strip()
        if not text:
            continue
        if _is_stop_heading(para):
            break
        entries.append(text)

    return entries


_REF_NUMBER_PREFIX_RE = re.compile(r"^\[(\d+)\]")


def _check_ref_numbering(entries: list[str]) -> bool:
    """Check whether reference entries have [N] numbering prefix."""
    if not entries:
        return False
    numbered = sum(1 for e in entries if _REF_NUMBER_PREFIX_RE.match(e))
    return numbered >= len(entries) * 0.5


# ---------------------------------------------------------------------------
# Phase 1: Style determination
# ---------------------------------------------------------------------------

def _determine_style(
    numbered_hits: int,
    author_year_count: int,
    note_count: int,
) -> Literal["numbered", "author_year", "footnote", "mixed", "none"]:
    """Determine the dominant citation style from signal counts."""
    signals = {
        "numbered": numbered_hits,
        "author_year": author_year_count,
        "footnote": note_count,
    }
    total = sum(signals.values())
    if total == 0:
        return "none"

    # Sort by count descending
    ranked = sorted(signals.items(), key=lambda x: -x[1])
    top_style, top_count = ranked[0]
    second_style, second_count = ranked[1]

    if top_count == 0:
        return "none"

    # If two styles both have significant presence, it's mixed
    if second_count >= 3 and second_count >= top_count * 0.3:
        # Special case: numbered + author-year often coexist in Chinese papers
        # where authors write "(张三，2020)[1]" — numbered is the real style
        if top_style == "numbered" and second_style == "author_year":
            return "numbered"
        return "mixed"

    return top_style  # type: ignore[return-value]


# ---------------------------------------------------------------------------
# Phase 2: Universal checkers
# ---------------------------------------------------------------------------

def _check_no_reference_section(
    doc: Document, profile: CitationProfile
) -> dict[str, Any] | None:
    """Flag when inline citations exist but no reference section is found."""
    if profile.has_ref_section:
        return None
    total_cites = profile.numbered_hit_count + profile.inline_author_year_count
    if total_cites == 0:
        return None
    return {
        "code": "CITATION_NO_REF_SECTION",
        "severity": "warning",
        "message": (
            f"Found {total_cites} inline citation(s) "
            f"but no reference section (参考文献/References)"
        ),
        "count": total_cites,
        "fixable": False,
    }


def _check_style_mixed(
    doc: Document, profile: CitationProfile
) -> dict[str, Any] | None:
    """Flag mixed citation styles."""
    if profile.style != "mixed":
        return None
    return {
        "code": "CITATION_STYLE_MIXED",
        "severity": "warning",
        "message": "Document mixes multiple citation styles (numbered + author-year)",
        "count": 1,
        "fixable": False,
    }


# ---------------------------------------------------------------------------
# Phase 2: Numbered style checkers
# ---------------------------------------------------------------------------

def _check_ref_numbering_missing(
    doc: Document, profile: CitationProfile
) -> dict[str, Any] | None:
    """Reference entries lack [N] numbering prefix."""
    if profile.ref_has_numbering:
        return None
    if profile.ref_entry_count == 0:
        return None
    return {
        "code": "CITATION_REF_NUMBERING_MISSING",
        "severity": "warning",
        "message": (
            f"Reference list has {profile.ref_entry_count} entries "
            f"but most lack [N] numbering prefix"
        ),
        "count": profile.ref_entry_count,
    }


def _check_number_out_of_range(
    doc: Document, profile: CitationProfile
) -> dict[str, Any] | None:
    """Inline [N] where N exceeds reference count."""
    if profile.ref_entry_count == 0:
        return None
    out_of_range = {
        n for n in profile.inline_numbers
        if n > profile.ref_entry_count
    }
    if not out_of_range:
        return None
    return {
        "code": "CITATION_OUT_OF_RANGE",
        "severity": "warning",
        "message": (
            f"{len(out_of_range)} citation(s) reference numbers "
            f"beyond the {profile.ref_entry_count} reference entries: "
            f"{sorted(out_of_range)}"
        ),
        "count": len(out_of_range),
        "fixable": False,
        "params": {"numbers": sorted(out_of_range), "ref_count": profile.ref_entry_count},
    }


def _check_number_gaps(
    doc: Document, profile: CitationProfile
) -> dict[str, Any] | None:
    """Inline citation numbers are not contiguous (e.g., 1,2,5 — missing 3,4)."""
    if not profile.inline_numbers:
        return None
    max_num = max(profile.inline_numbers)
    expected = set(range(1, max_num + 1))
    gaps = expected - profile.inline_numbers
    if not gaps:
        return None
    return {
        "code": "CITATION_GAP",
        "severity": "info",
        "message": (
            f"Citation numbers skip {len(gaps)} value(s): "
            f"{sorted(gaps)}"
        ),
        "count": len(gaps),
        "fixable": False,
        "params": {"missing": sorted(gaps)},
    }


def _check_never_cited(
    doc: Document, profile: CitationProfile
) -> dict[str, Any] | None:
    """Reference entries that are never cited in body text."""
    if profile.ref_entry_count == 0 or not profile.inline_numbers:
        return None
    ref_nums = set(range(1, profile.ref_entry_count + 1))
    never_cited = ref_nums - profile.inline_numbers
    if not never_cited:
        return None

    # Build entries map: number -> truncated text (for frontend display)
    entries_map: dict[int, str] = {}
    for n in sorted(never_cited):
        idx = n - 1  # 0-based
        if 0 <= idx < len(profile.ref_entries):
            text = profile.ref_entries[idx]
            entries_map[n] = text[:100]

    return {
        "code": "CITATION_NEVER_CITED",
        "severity": "info",
        "message": (
            f"{len(never_cited)} reference(s) never cited in body text: "
            f"{sorted(never_cited)}"
        ),
        "count": len(never_cited),
        "fixable": False,
        "params": {"numbers": sorted(never_cited), "entries": entries_map},
    }


# ---------------------------------------------------------------------------
# Checker registries
# ---------------------------------------------------------------------------

_UNIVERSAL_CHECKERS = [
    _check_no_reference_section,
    _check_style_mixed,
]

_STYLE_CHECKERS: dict[str, list] = {
    "numbered": [
        _check_ref_numbering_missing,
        _check_number_out_of_range,
        _check_number_gaps,
        _check_never_cited,
    ],
    "author_year": [],  # future: _check_author_not_found, _check_ref_orphan
    "footnote": [],     # future: _check_footnote_empty
    "mixed": [
        # Still run numbered checks if we have numbered data
        _check_ref_numbering_missing,
        _check_number_out_of_range,
        _check_number_gaps,
        _check_never_cited,
    ],
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _parse_number_list(raw: str) -> list[int]:
    """Parse '1', '1,3', '1-3', '1,3-5' into list of ints."""
    numbers: list[int] = []
    parts = re.split(r"[,，]", raw)
    for part in parts:
        part = part.strip()
        if "-" in part:
            try:
                a, b = part.split("-", 1)
                for n in range(int(a.strip()), int(b.strip()) + 1):
                    numbers.append(n)
            except (ValueError, TypeError):
                pass
        else:
            try:
                numbers.append(int(part))
            except (ValueError, TypeError):
                pass
    return numbers
