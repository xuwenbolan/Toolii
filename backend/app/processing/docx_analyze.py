"""DOCX document analysis: metadata extraction and health issue detection.

Checks are based on the thesis format standard defined in
docs/references/thesis-format-standard.md (Chinese undergraduate thesis).
"""

from __future__ import annotations

import io
import re
from collections import Counter
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Emu


# ---------------------------------------------------------------------------
# Thesis format standard values (from thesis-format-standard.md)
# ---------------------------------------------------------------------------

# Fonts
STD_BODY_CJK_FONT = "宋体"  # SimSun
STD_BODY_LATIN_FONT = "Times New Roman"
STD_HEADING_CJK_FONT = "宋体"  # SimSun for H1/H2/H3
STD_CAPTION_FONT = "黑体"  # SimHei

# Font sizes (in pt)
STD_BODY_SIZE_PT = 12.0  # 小四
STD_H1_SIZE_PT = 15.0  # 小三
STD_H2_SIZE_PT = 14.0  # 四号
STD_H3_SIZE_PT = 12.0  # 小四
STD_REF_SIZE_PT = 10.5  # 五号

# Spacing
STD_BODY_LINE_SPACING = 1.5  # 1.5x
STD_BODY_INDENT_CHARS = 200  # firstLineChars=200 (2 chars)
STD_BODY_INDENT_PT = 24.0  # ~2 chars at 12pt

# Page layout (twips)
STD_MARGIN_LR_TWIPS = 1803  # 3.18cm
STD_MARGIN_TB_TWIPS = 1440  # 2.54cm
STD_MARGIN_TOLERANCE = 60  # tolerance in twips (~1mm)
STD_A4_WIDTH = 11906
STD_A4_HEIGHT = 16838

# Heading numbering patterns (from standard Section 3.8)
_CN_LEVEL_PATTERNS = [
    (1, re.compile(r"^[一二三四五六七八九十百]+、")),
    (2, re.compile(r"^[（(]\s*[一二三四五六七八九十百]+\s*[）)]")),
    (3, re.compile(r"^\d+\.\s")),
    (3, re.compile(r"^\d+、")),
]


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def analyze_docx(docx_bytes: bytes) -> dict[str, Any]:
    """Parse DOCX and return analysis data as a plain dict."""
    doc = Document(io.BytesIO(docx_bytes))

    metadata = _extract_metadata(doc)
    headings = _extract_headings(doc)
    issues = _run_detectors(doc)

    heading_issue_codes = _heading_issue_map(issues)
    for h in headings:
        key = (h["level"], h["text"])
        if key in heading_issue_codes:
            h["has_issue"] = True
            h["issue_code"] = heading_issue_codes[key]

    score = _compute_score(issues)

    return {
        "metadata": metadata,
        "headings": headings,
        "issues": issues,
        "score": score,
    }


# ---------------------------------------------------------------------------
# Metadata extraction
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


def _extract_metadata(doc: Document) -> dict[str, Any]:
    word_count = 0
    paragraph_count = len(doc.paragraphs)
    heading_count = 0
    fonts: set[str] = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            word_count += len(text.split())
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            heading_count += 1
        for run in para.runs:
            if run.font.name:
                fonts.add(run.font.name)
            # Also collect CJK font from eastAsia attribute
            ea = _get_eastasia_font(run)
            if ea:
                fonts.add(ea)

    image_count = sum(1 for shape in doc.inline_shapes if shape.type is not None)
    style_count = len([s for s in doc.styles if s.name])
    page_estimate = _estimate_page_count(doc, word_count)

    return {
        "word_count": word_count,
        "paragraph_count": paragraph_count,
        "heading_count": heading_count,
        "image_count": image_count,
        "font_families": sorted(fonts),
        "style_count": style_count,
        "page_count_estimate": page_estimate,
    }


def _estimate_page_count(doc: Document, word_count: int) -> int:
    try:
        app_part = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties"
        )
        if app_part is not None:
            from lxml import etree
            tree = etree.fromstring(app_part.blob)
            ns = {"ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}
            pages_el = tree.find("ep:Pages", ns)
            if pages_el is not None and pages_el.text:
                count = int(pages_el.text)
                if count > 0:
                    return count
    except Exception:
        pass
    return max(1, round(word_count / 250))


def _extract_headings(doc: Document) -> list[dict[str, Any]]:
    headings: list[dict[str, Any]] = []
    for para in doc.paragraphs:
        if not para.style or not para.style.name:
            continue
        m = _HEADING_RE.match(para.style.name)
        if m:
            headings.append({
                "level": int(m.group(1)),
                "text": para.text.strip() or "(empty heading)",
                "has_issue": False,
                "issue_code": None,
            })
    return headings


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _is_heading(para) -> bool:
    return bool(para.style and para.style.name and para.style.name.startswith("Heading"))


def _get_body_paras(doc: Document):
    """Yield non-empty, non-heading paragraphs (body text candidates)."""
    for para in doc.paragraphs:
        if _is_heading(para):
            continue
        if not para.text.strip():
            continue
        yield para


def _get_heading_level(para) -> int | None:
    if not para.style or not para.style.name:
        return None
    m = _HEADING_RE.match(para.style.name)
    return int(m.group(1)) if m else None


def _get_run_font_name(run) -> str | None:
    """Get the effective font name for a run, checking rFonts XML for CJK."""
    # python-docx run.font.name returns ascii/hAnsi font
    return run.font.name


def _get_eastasia_font(run) -> str | None:
    """Read the eastAsia font from the run's XML rFonts element."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:eastAsia"))


def _get_line_spacing_xml(para) -> tuple[int | None, str | None]:
    """Read line spacing value and rule from paragraph XML.

    Returns (line_value, line_rule) where line_value is in 240ths of a line
    for 'auto' rule, or twips for 'exact'/'atLeast'.
    """
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        return None, None
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        return None, None
    line = spacing.get(qn("w:line"))
    rule = spacing.get(qn("w:lineRule"))
    return (int(line) if line else None), rule


def _get_first_line_indent_chars(para) -> int | None:
    """Read firstLineChars from paragraph XML (in hundredths of a character)."""
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        return None
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        return None
    val = ind.get(qn("w:firstLineChars"))
    return int(val) if val else None


# ---------------------------------------------------------------------------
# Thesis-standard detectors
# ---------------------------------------------------------------------------

def _check_body_font(doc: Document) -> list[dict[str, Any]]:
    """FT-001/002: Check body text fonts against standard (SimSun + TNR)."""
    wrong_latin: Counter[str] = Counter()
    wrong_cjk: Counter[str] = Counter()
    total_runs = 0

    for para in _get_body_paras(doc):
        for run in para.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            # Check Latin font (run.font.name)
            latin = run.font.name
            if latin and latin != STD_BODY_LATIN_FONT:
                wrong_latin[latin] += 1
            # Check CJK font (eastAsia attribute)
            cjk = _get_eastasia_font(run)
            if cjk and cjk != STD_BODY_CJK_FONT:
                wrong_cjk[cjk] += 1

    issues = []
    if wrong_latin:
        top = wrong_latin.most_common(3)
        issues.append({
            "code": "BODY_FONT_WRONG",
            "severity": "critical",
            "message": (
                f"Body Latin font should be {STD_BODY_LATIN_FONT}; "
                f"found {', '.join(f'{f}({c})' for f, c in top)}"
            ),
            "count": sum(wrong_latin.values()),
            "params": {"expected": STD_BODY_LATIN_FONT, "found": dict(top)},
        })
    if wrong_cjk:
        top = wrong_cjk.most_common(3)
        issues.append({
            "code": "BODY_CJK_FONT_WRONG",
            "severity": "critical",
            "message": (
                f"Body CJK font should be {STD_BODY_CJK_FONT}; "
                f"found {', '.join(f'{f}({c})' for f, c in top)}"
            ),
            "count": sum(wrong_cjk.values()),
            "params": {"expected": STD_BODY_CJK_FONT, "found": dict(top)},
        })
    return issues


def _check_body_size(doc: Document) -> dict[str, Any] | None:
    """FS-001: Check body text font size against standard (12pt / 小四)."""
    wrong_count = 0
    wrong_sizes: Counter[float] = Counter()
    total = 0

    for para in _get_body_paras(doc):
        for run in para.runs:
            if not run.text.strip() or not run.font.size:
                continue
            total += 1
            pt = run.font.size.pt
            if abs(pt - STD_BODY_SIZE_PT) > 0.5:
                wrong_count += 1
                wrong_sizes[pt] += 1

    if wrong_count > 0 and total > 0:
        top = wrong_sizes.most_common(3)
        return {
            "code": "BODY_SIZE_WRONG",
            "severity": "critical",
            "message": (
                f"Body text should be {STD_BODY_SIZE_PT}pt (小四); "
                f"{wrong_count}/{total} runs use other sizes"
            ),
            "count": wrong_count,
            "params": {"expected_pt": STD_BODY_SIZE_PT, "found": dict(top)},
        }
    return None


def _check_body_spacing(doc: Document) -> dict[str, Any] | None:
    """SP-001: Check body line spacing against standard (1.5x)."""
    wrong_count = 0
    total = 0

    for para in _get_body_paras(doc):
        total += 1
        line_val, rule = _get_line_spacing_xml(para)
        # 1.5x line spacing = line=360 with lineRule=auto (or no rule)
        if line_val is not None:
            if rule in (None, "auto"):
                if line_val != 360:
                    wrong_count += 1
            else:
                # exact or atLeast spacing — not 1.5x auto
                wrong_count += 1
        # None means inherited from style — assume OK unless style is wrong

    if wrong_count > 0 and total > 0:
        return {
            "code": "BODY_SPACING_WRONG",
            "severity": "critical",
            "message": f"Body line spacing should be 1.5x; {wrong_count}/{total} paragraphs differ",
            "count": wrong_count,
        }
    return None


def _check_body_indent(doc: Document) -> dict[str, Any] | None:
    """SP-002: Check body first-line indent against standard (2 chars)."""
    wrong_count = 0
    total = 0

    for para in _get_body_paras(doc):
        total += 1
        # Check firstLineChars in XML (most reliable)
        chars = _get_first_line_indent_chars(para)
        if chars is not None:
            if chars != STD_BODY_INDENT_CHARS:
                wrong_count += 1
        else:
            # Fallback: check EMU-based indent
            fi = para.paragraph_format.first_line_indent
            if fi is not None:
                pt = fi / 12700  # EMU to pt
                if abs(pt - STD_BODY_INDENT_PT) > 4:  # tolerance ~4pt
                    wrong_count += 1
            else:
                # No indent set at all
                wrong_count += 1

    if wrong_count > 0 and total > 0:
        return {
            "code": "BODY_INDENT_WRONG",
            "severity": "critical",
            "message": f"Body first-line indent should be 2 chars; {wrong_count}/{total} paragraphs differ",
            "count": wrong_count,
        }
    return None


def _check_body_alignment(doc: Document) -> dict[str, Any] | None:
    """SP-007: Check body text alignment against standard (justify)."""
    wrong_count = 0
    total = 0

    for para in _get_body_paras(doc):
        total += 1
        align = para.paragraph_format.alignment
        if align is not None and align != WD_ALIGN_PARAGRAPH.JUSTIFY:
            wrong_count += 1

    if wrong_count > 0 and total > 0:
        return {
            "code": "BODY_ALIGNMENT_WRONG",
            "severity": "warning",
            "message": f"Body text should be justify-aligned; {wrong_count}/{total} paragraphs differ",
            "count": wrong_count,
        }
    return None


def _check_heading_size(doc: Document) -> list[dict[str, Any]]:
    """FS-002/003/004: Check heading font sizes against standard."""
    expected = {1: STD_H1_SIZE_PT, 2: STD_H2_SIZE_PT, 3: STD_H3_SIZE_PT}
    level_wrong: dict[int, int] = {}

    for para in doc.paragraphs:
        level = _get_heading_level(para)
        if level is None or level not in expected:
            continue
        runs = [r for r in para.runs if r.text.strip() and r.font.size]
        if not runs:
            continue
        # Check dominant size of the heading
        for run in runs:
            pt = run.font.size.pt
            if abs(pt - expected[level]) > 0.5:
                level_wrong[level] = level_wrong.get(level, 0) + 1
                break

    issues = []
    cn_size_names = {1: "小三", 2: "四号", 3: "小四"}
    for level, count in sorted(level_wrong.items()):
        issues.append({
            "code": "HEADING_SIZE_WRONG",
            "severity": "critical" if level == 1 else "warning",
            "message": (
                f"Heading {level} should be {expected[level]}pt ({cn_size_names[level]}); "
                f"{count} heading(s) differ"
            ),
            "count": count,
            "params": {"level": level, "expected_pt": expected[level]},
        })
    return issues


def _check_heading_alignment(doc: Document) -> dict[str, Any] | None:
    """SP-003: Check H1 alignment against standard (center)."""
    wrong_count = 0
    total = 0

    for para in doc.paragraphs:
        level = _get_heading_level(para)
        if level != 1:
            continue
        total += 1
        align = para.paragraph_format.alignment
        if align is not None and align != WD_ALIGN_PARAGRAPH.CENTER:
            wrong_count += 1

    if wrong_count > 0 and total > 0:
        return {
            "code": "HEADING_ALIGNMENT_WRONG",
            "severity": "warning",
            "message": f"Heading 1 should be center-aligned; {wrong_count}/{total} heading(s) differ",
            "count": wrong_count,
        }
    return None


def _check_page_margins(doc: Document) -> dict[str, Any] | None:
    """PL-002/003: Check page margins against standard."""
    try:
        section = doc.sections[0]
    except (IndexError, TypeError):
        return None

    problems = []
    tol = STD_MARGIN_TOLERANCE

    if section.left_margin and abs(section.left_margin - STD_MARGIN_LR_TWIPS * 635) > tol * 635:
        problems.append("left")
    if section.right_margin and abs(section.right_margin - STD_MARGIN_LR_TWIPS * 635) > tol * 635:
        problems.append("right")
    if section.top_margin and abs(section.top_margin - STD_MARGIN_TB_TWIPS * 635) > tol * 635:
        problems.append("top")
    if section.bottom_margin and abs(section.bottom_margin - STD_MARGIN_TB_TWIPS * 635) > tol * 635:
        problems.append("bottom")

    if problems:
        return {
            "code": "PAGE_MARGINS_WRONG",
            "severity": "warning",
            "message": f"Page margins ({', '.join(problems)}) differ from standard (LR 3.18cm, TB 2.54cm)",
            "count": len(problems),
            "params": {"wrong_sides": problems},
        }
    return None


# ---------------------------------------------------------------------------
# Generic detectors (kept from previous implementation)
# ---------------------------------------------------------------------------

def _check_toc_outdated(doc: Document) -> dict[str, Any] | None:
    """Check if document has a TOC that may need updating.

    Skips if the TOC fields are already marked dirty (i.e., repair was applied
    and the TOC will refresh on next open in Word/LibreOffice).
    """
    body = doc.element.body

    # Check SDT-wrapped TOC
    for sdt in body.iter(qn("w:sdt")):
        sdt_pr = sdt.find(qn("w:sdtPr"))
        if sdt_pr is not None:
            doc_part = sdt_pr.find(qn("w:docPartObj"))
            if doc_part is not None:
                gallery = doc_part.find(qn("w:docPartGallery"))
                if gallery is not None and "Table of Contents" in (gallery.get(qn("w:val")) or ""):
                    # Check if already marked dirty
                    if _has_dirty_fields(sdt):
                        return None  # Already marked for refresh
                    return {"code": "TOC_OUTDATED", "severity": "warning",
                            "message": "Document contains a TOC that may be outdated", "count": 1}

    # Check bare TOC field codes
    for fld in body.iter(qn("w:fldChar")):
        if fld.get(qn("w:fldCharType")) != "begin":
            continue
        if fld.get(qn("w:dirty")) == "true":
            return None  # Already marked dirty
        parent = fld.getparent()
        if parent is None:
            continue
        for sibling in parent.itersiblings():
            instr = sibling.find(qn("w:instrText"))
            if instr is not None and instr.text and "TOC" in instr.text:
                return {"code": "TOC_OUTDATED", "severity": "warning",
                        "message": "Document contains a TOC that may be outdated", "count": 1}
            break
    return None


def _has_dirty_fields(element) -> bool:
    """Check if any field characters within element are marked dirty."""
    for fld in element.iter(qn("w:fldChar")):
        if fld.get(qn("w:dirty")) == "true":
            return True
    return False


def _check_redundant_empty_paragraphs(doc: Document) -> dict[str, Any] | None:
    consecutive = 0
    total_redundant = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            consecutive += 1
            if consecutive >= 2:
                total_redundant += 1
        else:
            consecutive = 0
    if total_redundant > 0:
        return {"code": "REDUNDANT_EMPTY_PARAGRAPHS", "severity": "info",
                "message": f"Found {total_redundant} redundant empty paragraphs",
                "count": total_redundant}
    return None


def _check_missing_heading_styles(doc: Document) -> dict[str, Any] | None:
    """Detect paragraphs that look like headings but lack Heading styles."""
    suspects: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) > 100 or _is_heading(para):
            continue
        runs = [r for r in para.runs if r.text.strip()]
        if not runs:
            continue
        all_bold = all(r.bold for r in runs)
        any_large = any(r.font.size and r.font.size >= Pt(14) for r in runs)
        if all_bold and any_large:
            suspects.append(text[:50])

    if suspects:
        return {"code": "MISSING_HEADING_STYLES", "severity": "warning",
                "message": f"{len(suspects)} paragraph(s) look like headings but lack Heading styles",
                "count": len(suspects), "params": {"examples": suspects[:5]}}
    return None


def _check_heading_orphan(doc: Document) -> dict[str, Any] | None:
    missing_count = 0
    for para in doc.paragraphs:
        if not _is_heading(para):
            continue
        if not para.paragraph_format.keep_with_next:
            missing_count += 1
    if missing_count > 0:
        return {"code": "HEADING_ORPHAN", "severity": "warning",
                "message": f"{missing_count} heading(s) may appear orphaned at page bottom",
                "count": missing_count}
    return None


def _check_empty_heading(doc: Document) -> dict[str, Any] | None:
    empty_count = 0
    for para in doc.paragraphs:
        if _is_heading(para) and not para.text.strip():
            empty_count += 1
    if empty_count > 0:
        return {"code": "EMPTY_HEADING", "severity": "warning",
                "message": f"{empty_count} heading(s) have no text content",
                "count": empty_count, "fixable": False}
    return None


def _check_heading_level_flat(doc: Document) -> dict[str, Any] | None:
    """Detect all headings at same level when numbering suggests sub-levels."""
    levels, texts = [], []
    for para in doc.paragraphs:
        level = _get_heading_level(para)
        if level is not None:
            levels.append(level)
            texts.append(para.text.strip())
    if len(levels) < 3 or len(set(levels)) != 1:
        return None
    inferred: dict[int, int] = {}
    for text in texts:
        for expected_level, pattern in _CN_LEVEL_PATTERNS:
            if pattern.match(text):
                inferred[expected_level] = inferred.get(expected_level, 0) + 1
                break
    if len(inferred) > 1:
        return {"code": "HEADING_LEVEL_FLAT", "severity": "warning",
                "message": f"All {len(levels)} headings are H{levels[0]}, but numbering suggests multiple levels",
                "count": len(levels),
                "params": {"current_level": levels[0], "inferred_levels": inferred}}
    return None


def _check_heading_format_inconsistency(doc: Document) -> dict[str, Any] | None:
    """Detect same-level headings with inconsistent direct formatting."""
    # Standard sizes for comparison (None = inherited, treat as equivalent to standard)
    std_sizes = {1: STD_H1_SIZE_PT, 2: STD_H2_SIZE_PT, 3: STD_H3_SIZE_PT}

    level_formats: dict[int, list[tuple]] = {}
    for para in doc.paragraphs:
        level = _get_heading_level(para)
        if level is None:
            continue
        runs = [r for r in para.runs if r.text.strip()]
        if not runs:
            continue
        r = runs[0]
        # Normalize bold: None (inherited) and True are both "bold" for headings
        bold = True if r.bold in (True, None) else False
        # Normalize size: None (inherited) treated as standard size for that level
        size = float(r.font.size.pt) if r.font.size else None
        std = std_sizes.get(level)
        if size is None or (std and abs(size - std) < 0.5):
            size = std  # Treat inherited or standard-matching as "standard"
        fmt = (r.font.name, size, bold)
        level_formats.setdefault(level, []).append(fmt)

    bad = [lv for lv, fmts in level_formats.items() if len(fmts) >= 2 and len(set(fmts)) > 1]
    if bad:
        return {"code": "HEADING_FORMAT_INCONSISTENCY", "severity": "warning",
                "message": f"Heading level(s) {', '.join(f'H{l}' for l in bad)} have inconsistent formatting",
                "count": len(bad), "params": {"levels": bad}}
    return None


def _check_image_overflow(doc: Document) -> dict[str, Any] | None:
    try:
        section = doc.sections[0]
        pw = section.page_width - section.left_margin - section.right_margin
    except (IndexError, TypeError):
        return None
    count = sum(1 for s in doc.inline_shapes if s.width and s.width > pw)
    if count > 0:
        return {"code": "IMAGE_OVERFLOW", "severity": "warning",
                "message": f"{count} image(s) exceed page width", "count": count}
    return None


def _check_numbering_discontinuity(doc: Document) -> dict[str, Any] | None:
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return None
    except Exception:
        return None
    num_ids: set[str] = set()
    for para in doc.paragraphs:
        p_pr = para._element.find(qn("w:pPr"))
        if p_pr is None:
            continue
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            el = num_pr.find(qn("w:numId"))
            if el is not None and el.get(qn("w:val")):
                num_ids.add(el.get(qn("w:val")))
    defined = set()
    for num in numbering_part._element.iter(qn("w:num")):
        nid = num.get(qn("w:numId"))
        if nid:
            defined.add(nid)
    orphan = num_ids - defined - {"0"}
    if orphan:
        return {"code": "NUMBERING_DISCONTINUITY", "severity": "warning",
                "message": f"{len(orphan)} list numbering reference(s) point to missing definitions",
                "count": len(orphan), "fixable": False}
    return None


def _check_heading_hierarchy_gap(doc: Document) -> dict[str, Any] | None:
    levels = [_get_heading_level(p) for p in doc.paragraphs if _get_heading_level(p) is not None]
    gaps = [f"H{levels[i-1]}->H{levels[i]}" for i in range(1, len(levels)) if levels[i] > levels[i-1]+1]
    if gaps:
        return {"code": "HEADING_HIERARCHY_GAP", "severity": "info",
                "message": f"Heading level gap(s): {', '.join(gaps[:5])}",
                "count": len(gaps), "fixable": False}
    return None


def _check_broken_cross_references(doc: Document) -> dict[str, Any] | None:
    body = doc.element.body
    bookmarks = {bm.get(qn("w:name")) for bm in body.iter(qn("w:bookmarkStart")) if bm.get(qn("w:name"))}
    ref_targets: Counter[str] = Counter()
    for instr in body.iter(qn("w:instrText")):
        if instr.text and "REF " in instr.text:
            m = re.search(r"REF\s+(\S+)", instr.text)
            if m:
                ref_targets[m.group(1)] += 1
    broken = {t: c for t, c in ref_targets.items() if t not in bookmarks}
    if broken:
        total = sum(broken.values())
        return {"code": "BROKEN_CROSS_REFERENCES", "severity": "warning",
                "message": f"{total} cross-reference(s) point to {len(broken)} missing bookmark(s): {list(broken.keys())[:5]}",
                "count": total, "fixable": False, "params": {"broken": broken}}
    return None


def _check_style_bloat(doc: Document) -> dict[str, Any] | None:
    all_s = {s.name for s in doc.styles if s.name}
    used = {p.style.name for p in doc.paragraphs if p.style and p.style.name}
    unused = len(all_s - used)
    if unused > 20:
        return {"code": "STYLE_BLOAT", "severity": "info",
                "message": f"{unused} unused styles out of {len(all_s)} total",
                "count": unused, "fixable": False}
    return None


def _check_citations(doc: Document) -> list[dict[str, Any]]:
    from app.processing.docx_citations import check_citations
    return check_citations(doc)


# ---------------------------------------------------------------------------
# Detector registry
# ---------------------------------------------------------------------------

_DETECTORS = [
    # Thesis-standard checks (return list or single)
    _check_body_font,       # FT-001/002
    _check_body_size,       # FS-001
    _check_body_spacing,    # SP-001
    _check_body_indent,     # SP-002
    _check_body_alignment,  # SP-007
    _check_heading_size,    # FS-002/003/004
    _check_heading_alignment,  # SP-003
    _check_page_margins,    # PL-002/003
    # Generic structural checks
    _check_toc_outdated,
    _check_redundant_empty_paragraphs,
    _check_missing_heading_styles,
    _check_heading_orphan,
    _check_empty_heading,
    _check_heading_level_flat,
    _check_heading_format_inconsistency,
    _check_image_overflow,
    _check_numbering_discontinuity,
    # Citations
    _check_citations,
    # Detect-only
    _check_heading_hierarchy_gap,
    _check_broken_cross_references,
    _check_style_bloat,
]


def _run_detectors(doc: Document) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for detector in _DETECTORS:
        result = detector(doc)
        if result is None:
            continue
        if isinstance(result, list):
            issues.extend(result)
        else:
            issues.append(result)
    return issues


# ---------------------------------------------------------------------------
# Scoring
# ---------------------------------------------------------------------------

_SEVERITY_PENALTY = {"critical": 20, "warning": 10, "info": 5}


def _compute_score(issues: list[dict[str, Any]]) -> int:
    """Compute health score (0-100).

    Only fixable issues reduce the score. Unfixable (detect-only) issues
    are informational — they are displayed to the user but do not affect
    the score, since the user cannot resolve them through our repair tool.
    """
    score = 100
    for issue in issues:
        if not issue.get("fixable", True):
            continue
        score -= _SEVERITY_PENALTY.get(issue.get("severity", "info"), 5)
    return max(0, score)


def _heading_issue_map(issues: list[dict[str, Any]]) -> dict[tuple[int, str], str]:
    return {}
